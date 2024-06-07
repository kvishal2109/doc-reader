import os
import openai
import json
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.faiss import FAISS
from langchain.document_loaders import TextLoader
from langchain_community.document_loaders import Docx2txtLoader
from get_config import Config
from docx import Document
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import ChatPromptTemplate, PromptTemplate

# from langchain.document_loaders.parsers.pdf import PDFPlumberParser

from langchain_community.document_loaders import AmazonTextractPDFLoader
import tabula

from flask import Flask, request, jsonify
from flask_cors import CORS
import os

app = Flask(__name__)
CORS(app)

def data_pre_processing(file_path):

    print("file_path:::::::::::",file_path)
    splited_file_path = file_path.split("/")
    splited_file_path = splited_file_path[len(splited_file_path)-1]
    splited_file_path = splited_file_path.split(".")
    print("splited_file_path:::::::::::::::",splited_file_path)
    file_name = splited_file_path[0]
    file_ext = splited_file_path[1]

    print("file_name:::::::::::::::",file_name)
    print("file_ext:::::::::::::::",file_ext)


    if(file_ext == "docx"):
        doc_docx = Document(file_path)

        with open(f"./docs/{file_name}.txt", 'w') as txt_file:
            for paragraph in doc_docx.paragraphs:
                txt_file.write(paragraph.text + '\n')

        # Extract tables
        with open(f"./docs/{file_name}.txt", 'a') as txt_file:  # Append mode to add tables after text
            for table in doc_docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt_file.write(cell.text + '\t')
                    txt_file.write('\n')
                txt_file.write('\n')

    elif(file_ext == "pdf"):
        print("fileext::::::::::::",file_ext)
        print("file_path::::::::::::",file_path)
        tables_pdf = tabula.read_pdf(file_path, pages='all')

        print("tables_pdf::::::::::::::::::::::::",tables_pdf)
        with open(f"./docs/{file_name}.txt", 'w') as txt_file:
            for table in tables_pdf:
                txt_file.write(str(table))
                txt_file.write('\n\n')

    os.remove(file_path)


@app.route('/api/v1/pre/processing', methods=['POST'])
def pre_process_file():
    files = request.files.getlist('file')

    for file in files:
        file_path = "./docs/" + file.filename
        file.save(file_path)

        data_pre_processing(file_path)
        

    return jsonify({
        "msg": "got the file"++++++++++++++++++++++++++++++++++++++++++
    }), 200

if _name_ == '__main__':
    app.run(debug=True,port='5500')
